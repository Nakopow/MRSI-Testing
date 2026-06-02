import glob
import html
import json
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from scraper import enrich_articles, fetch_rss_items, load_url_cache
from summarizer import generate_daily_digest

# Cache file path for URL data persistence
URL_CACHE_FILE = ".url_cache.json"
TOPIC_LABELS = {
    "ai": "Artificial Intelligence",
    "cybersecurity": "Cybersecurity",
    "web3": "Web3 / Blockchain",
}


def save_digest_as_docx(digest: str, filename: str = "Daily_Tech_Briefing.docx") -> None:
    """
    Save the generated digest as a formatted Word document.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    today_date = datetime.now().strftime("%B %d, %Y")

    title = doc.add_paragraph()
    title_run = title.add_run("MARKET RESEARCH AND STRATEGIC INSIGHT")
    title_run.bold = True
    title_run.font.size = Pt(16)

    date_paragraph = doc.add_paragraph(today_date)
    date_paragraph.runs[0].italic = True

    doc.add_paragraph()

    lines = digest.split("\n")

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            continue

        if line.startswith("# "):
            doc.add_heading(line.replace("# ", "").strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", "").strip(), level=2)
        elif "Briefing:" in line:
            doc.add_heading(line, level=2)
        elif line.lower() == "executive summary":
            doc.add_heading("Executive Summary", level=3)
        elif line.lower().startswith("sources"):
            doc.add_heading("Sources", level=3)
        elif "targeted questions" in line.lower():
            doc.add_heading(line, level=3)
        elif line.startswith("●") or line.startswith("•") or line.startswith("*"):
            clean_line = line.lstrip("●•* ").strip()
            doc.add_paragraph(clean_line, style="List Bullet")
        elif any(line.startswith(f"{i}.") for i in range(1, 10)):
            doc.add_paragraph(line, style="List Number")
        elif set(line) == {"="}:
            continue
        else:
            doc.add_paragraph(line)

    doc.add_paragraph()
    footer = doc.add_paragraph("Send your feedback to exoasia.joddgabrielvillegas@gmail.com")
    footer.runs[0].italic = True

    doc.save(filename)


def run_formatter(output_txt: str) -> None:
    """
    Runs formatter.py after the daily digest is ready.
    Reads output.txt and produces one branded Daily Insight .docx per topic
    in the insights/ folder.
    Always runs — raises on critical failures instead of silently skipping.
    """
    import os
    import formatter as daily_formatter

    script_dir = os.path.dirname(os.path.abspath(__file__))
    insights_dir = os.path.join(script_dir, "insights")
    template = daily_formatter.TEMPLATE
    date_str = datetime.now().strftime("%B %d, %Y")
    date_slug = datetime.now().strftime("%b%d")

    print("\n" + "=" * 60)
    print("STEP 2 — Daily Insight Formatter (formatter.py)")
    print("=" * 60)

    if not os.path.exists(template):
        raise FileNotFoundError(
            f"Daily Insight template not found at '{template}'. "
            "Place Exoasia_MRSI_DailyInsight_Mar9_AI.docx next to main.py and re-run."
        )

    os.makedirs(insights_dir, exist_ok=True)

    print(f"  Parsing {output_txt} ...")
    topics = daily_formatter.split_topics(output_txt)

    if not topics:
        raise ValueError(
            f"No topics found in '{output_txt}'. "
            "Ensure the digest was generated correctly before running the formatter."
        )

    print(f"  Found {len(topics)} topic(s): {', '.join(topics)}")
    success_count = 0
    for key, data in topics.items():
        print(f"  Generating {data['display']} ...")
        content = daily_formatter.parse_topic(data["lines"])
        filename = f"Exoasia_MRSI_DailyInsight_{date_slug}_{data['slug']}.docx"
        out_path = os.path.join(insights_dir, filename)
        try:
            ok = daily_formatter.generate_docx(
                content, data["display"], date_str, out_path, template
            )
            if ok:
                print(f"    OK  -> insights/{filename}")
                success_count += 1
            else:
                raise RuntimeError(f"generate_docx returned False for topic '{key}' — check formatter logic.")
        except PermissionError:
            raise PermissionError(
                f"Cannot write '{filename}' — close the file in Word and re-run."
            )

    print(f"\n  Daily Insight formatter done. {success_count}/{len(topics)} .docx file(s) written to insights/")


def run_tl_pipeline(output_txt: str, platforms: list = None, topic: str = None) -> None:
    """
    Runs the Thought Leadership pipeline after the daily digest is ready.

    Step 1 — tl_summarizer: reads output.txt, calls Gemini per topic,
              generates images via Hugging Face, saves tl_output_{topic}.json.
    Step 2 — tl_formatter: reads each JSON file, builds branded .docx files
              in the TLPs folder.

    Args:
        output_txt: Path to the daily digest file
        platforms: Optional list of platform names to generate (e.g., ['linkedin', 'instagram']).
                   If None, generates for all platforms.
        topic: Optional topic to generate for (e.g., 'ai', 'cybersecurity', 'web3').
               If None or 'all', generates for all topics.

    Always runs — raises on critical failures instead of silently skipping.
    """
    import os
    import tl_summarizer
    import tl_formatter

    script_dir = os.path.dirname(os.path.abspath(__file__))
    tlp_dir = os.path.join(script_dir, "TLPs")
    prompts_path = os.path.join(script_dir, "tl_prompts_config.json")
    template = tl_formatter.TEMPLATE
    date_slug = datetime.now().strftime("%b%d")

    # ── Step 1: Summarizer ────────────────────────────────────────────────────
    print("\n" + "=" * 60)
    print("STEP 3 — Thought Leadership Summarizer (tl_summarizer.py)")
    print("=" * 60)

    gemini_key = os.environ.get("GEMINI_API_KEY")
    if not gemini_key:
        raise EnvironmentError(
            "GEMINI_API_KEY environment variable is not set. "
            "Export it before running: export GEMINI_API_KEY=your_key_here"
        )

    

    json_files = tl_summarizer.summarize_tl(
        gemini_key=gemini_key,
        output_txt=output_txt,
        prompts_path=prompts_path,
        outdir=script_dir,
    )

    if not json_files:
        raise RuntimeError(
            "TL summarizer produced no JSON files. "
            "Check that output.txt contains usable topic briefings and that the upstream daily digest "
            "did not write only 'Error: Failed to generate summary' placeholders."
        )

    # ── Step 2: Formatter ─────────────────────────────────────────────────────
    print("\n" + "=" * 60)
    print("STEP 4 — Thought Leadership Formatter (tl_formatter.py)")
    print("=" * 60)

    os.makedirs(tlp_dir, exist_ok=True)

    if not os.path.exists(template):
        raise FileNotFoundError(
            f"TL template not found at '{template}'. "
            "Place Exoasia_MRSI_ThoughtLeadership_Template.docx next to main.py and re-run."
        )

    success_count = 0
    for json_path in json_files:
        print(f"\n  Processing {os.path.basename(json_path)} ...")
        with open(json_path, encoding="utf-8") as f:
            data = json.load(f)

        topic_key = data.get("_topic_key", "unknown")
        slug = tl_formatter.TOPIC_SLUG.get(topic_key, topic_key.capitalize())
        out_path = os.path.join(tlp_dir, f"Exoasia_MRSI_ThoughtLeadership_{date_slug}_{slug}.docx")

        try:
            ok = tl_formatter.generate_docx(data, out_path, template=template)
            if ok:
                print(f"    OK  -> TLPs/{os.path.basename(out_path)}")
                success_count += 1
            else:
                raise RuntimeError(f"generate_docx returned False for '{os.path.basename(json_path)}'.")
        except PermissionError:
            raise PermissionError(
                f"Cannot write '{os.path.basename(out_path)}' — close the file in Word and re-run."
            )

    print(f"\n  TL formatter done. {success_count}/{len(json_files)} .docx file(s) written to TLPs/")


def normalize_text(value: str) -> str:
    """
    Repair common mojibake sequences and normalize punctuation for dashboard output.
    """
    if not value:
        return ""

    replacements = {
        "â€”": "—",
        "â€“": "–",
        "â€¦": "...",
        "â€™": "'",
        "â€œ": '"',
        "â€": '"',
        "â€¢": "•",
        "â—": "•",
        "Â·": "·",
        "Â": "",
    }
    for bad, good in replacements.items():
        value = value.replace(bad, good)
    return value.strip()


def html_text(value: str) -> str:
    return html.escape(normalize_text(value))


def slugify_topic(value: str) -> str:
    return (
        normalize_text(value).lower().replace("/", " ").replace("\\", " ").replace("  ", " ").strip()
        .replace(" ", "-")
    )


def format_timestamp(path: Path) -> str:
    return datetime.fromtimestamp(path.stat().st_mtime).strftime("%b %d, %Y %I:%M %p")


def parse_digest_sections(output_txt: str) -> Dict[str, dict]:
    """
    Parse output.txt into topic cards for the dashboard.
    """
    if not os.path.exists(output_txt):
        return {}

    with open(output_txt, "r", encoding="utf-8") as f:
        lines = [normalize_text(line.rstrip()) for line in f]

    sections: Dict[str, dict] = {}
    current = None
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("## ") and line[3:].strip().upper() in {
            "ARTIFICIAL INTELLIGENCE",
            "CYBERSECURITY",
            "WEB3 / BLOCKCHAIN",
        }:
            topic_name = line[3:].strip()
            key = "ai" if "ARTIFICIAL" in topic_name else "cybersecurity" if "CYBER" in topic_name else "web3"
            current = {
                "key": key,
                "label": topic_name.title() if key != "web3" else "Web3 / Blockchain",
                "title": "",
                "briefing_line": "",
                "summary": [],
                "details": [],
                "questions": [],
                "sources": [],
            }
            sections[key] = current
            i += 1
            continue

        if current:
            if not current["title"] and line and line not in {"Executive Summary", "Detailed Examination"} and not line.startswith("## "):
                next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
                if "Briefing:" in next_line:
                    current["title"] = line
                    current["briefing_line"] = next_line
            elif line.startswith("- "):
                current["summary"].append(line[2:].strip())
            elif line.startswith("Sources"):
                pass
            elif line.startswith("["):
                current["sources"].append(line)
            elif any(line.startswith(f"{n}.") for n in range(1, 6)):
                current["questions"].append(line)
            elif (
                line
                and "Briefing:" not in line
                and line not in {"Executive Summary", "Detailed Examination", "5 Questions for Thought Leadership"}
                and not line.startswith("## ")
                and set(line) != {"="}
            ):
                current["details"].append(line)
        i += 1

    return sections


def load_tlp_payloads() -> Dict[str, dict]:
    payloads: Dict[str, dict] = {}
    # Try real TLP output files first
    for path_str in glob.glob("tl_output_*.json"):
        path = Path(path_str)
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        topic_key = data.get("_topic_key") or path.stem.replace("tl_output_", "")
        payloads[topic_key] = data
    # Fall back to sample TLP files if no real ones found (e.g., on Vercel)
    if not payloads:
        for path_str in glob.glob("sample_tl_output_*.json"):
            path = Path(path_str)
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            topic_key = data.get("_topic_key") or path.stem.replace("sample_tl_output_", "")
            payloads[topic_key] = data
    return payloads


def extract_posting_windows(tlp_payloads: Dict[str, dict]) -> List[dict]:
    windows = []
    for topic_key, payload in tlp_payloads.items():
        for piece in payload.get("pieces", []):
            note = normalize_text(piece.get("posting_notes", "")).splitlines()[0] if piece.get("posting_notes") else "Posting notes available"
            windows.append(
                {
                    "topic": TOPIC_LABELS.get(topic_key, topic_key.title()),
                    "platform": normalize_text(piece.get("platform", "Platform")),
                    "note": note,
                }
            )
    return windows


def render_metric_card(label: str, value: str, tone: str = "") -> str:
    return f"""
        <div class="metric-card {tone}">
          <div class="metric-label">{html_text(label)}</div>
          <div class="metric-value">{html_text(value)}</div>
        </div>
    """


def render_insight_cards(digest_sections: Dict[str, dict], insight_files: List[Path]) -> str:
    cards = []
    for path in insight_files:
        topic_guess = "ai" if "AI" in path.name else "cybersecurity" if "Cyber" in path.name else "web3"
        section = digest_sections.get(topic_guess, {})
        bullets = "".join(
            f"<li>{html_text(item)}</li>" for item in section.get("summary", [])[:3]
        ) or "<li>No executive summary available yet.</li>"
        cards.append(
            f"""
            <article class="content-card">
              <div class="card-head">
                <div>
                  <div class="eyebrow">{html_text(TOPIC_LABELS.get(topic_guess, topic_guess.title()))}</div>
                  <h3>{html_text(section.get("title") or path.stem)}</h3>
                </div>
                <span class="doc-pill">DOCX</span>
              </div>
              <p class="muted">{html_text(section.get("briefing_line") or format_timestamp(path))}</p>
              <ul class="clean-list">{bullets}</ul>
              <div class="meta-row">
                <span>{html_text(path.name)}</span>
                <span>{html_text(format_timestamp(path))}</span>
              </div>
            </article>
            """
        )
    return "\n".join(cards) or '<div class="empty-state">No Daily Insight files found yet.</div>'


def render_tlp_cards(tlp_payloads: Dict[str, dict], tlp_files: List[Path]) -> str:
    cards = []
    file_lookup = {path.name.lower(): path for path in tlp_files}
    for topic_key, payload in tlp_payloads.items():
        heading = TOPIC_LABELS.get(topic_key, topic_key.title())
        pieces_html = []
        for piece in payload.get("pieces", []):
            body = piece.get("copy") or piece.get("script") or piece.get("caption") or ""
            excerpt = normalize_text(body).split("\n")[0]
            pieces_html.append(
                f"""
                <div class="platform-card">
                  <div class="platform-title">{html_text(piece.get("platform", "Platform"))}</div>
                  <div class="platform-angle">{html_text(piece.get("angle", ""))}</div>
                  <p>{html_text(excerpt[:240])}</p>
                </div>
                """
            )

        matching_file = next((path for path in tlp_files if topic_key.lower() in path.name.lower()), None)
        cards.append(
            f"""
            <section class="content-card">
              <div class="card-head">
                <div>
                  <div class="eyebrow">Thought Leadership Pack</div>
                  <h3>{html_text(heading)}</h3>
                </div>
                <span class="doc-pill accent">5 pieces</span>
              </div>
              <p class="muted">{html_text(normalize_text(payload.get("editorial_note", "")).splitlines()[0].lstrip("- "))}</p>
              <div class="platform-grid">
                {''.join(pieces_html)}
              </div>
              <div class="meta-row">
                <span>{html_text(matching_file.name if matching_file else 'No DOCX generated yet')}</span>
                <span>{html_text(format_timestamp(matching_file) if matching_file else payload.get('_date', ''))}</span>
              </div>
            </section>
            """
        )
    return "\n".join(cards) or '<div class="empty-state">No TLP JSON outputs found yet.</div>'


def render_autopost_page(posting_windows: List[dict]) -> str:
    items = []
    for window in posting_windows:
        items.append(
            f"""
            <div class="schedule-card">
              <div class="schedule-platform">{html_text(window['platform'])}</div>
              <div class="schedule-topic">{html_text(window['topic'])}</div>
              <p>{html_text(window['note'])}</p>
            </div>
            """
        )
    return "\n".join(items) or '<div class="empty-state">No posting windows available yet.</div>'


def render_schedule_page(digest_sections: Dict[str, dict], insight_files: List[Path], tlp_files: List[Path]) -> str:
    rows = []
    for key, section in digest_sections.items():
        rows.append(
            f"""
            <tr>
              <td>{html_text(TOPIC_LABELS.get(key, key.title()))}</td>
              <td>{html_text(section.get('briefing_line') or 'Generated from latest digest')}</td>
              <td>{len(section.get('sources', []))}</td>
              <td>{5 if tlp_files else 0}</td>
            </tr>
            """
        )
    return "\n".join(rows) or "<tr><td colspan='4'>No schedule data available yet.</td></tr>"


def render_api_status() -> str:
    keys = [
        ("Google Gemini", "GEMINI_API_KEY"),
        ("OpenAI", "OPENAI_API_KEY"),
        ("Mailchimp", "MAILCHIMP_API_KEY"),
    ]
    cards = []
    for label, env_name in keys:
        is_set = bool(os.environ.get(env_name))
        status = "Connected" if is_set else "Not connected"
        cards.append(
            f"""
            <div class="api-status {'ok' if is_set else 'off'}">
              <div class="api-name">{html_text(label)}</div>
              <div class="api-key">{html_text(env_name)}</div>
              <div class="api-pill">{html_text(status)}</div>
            </div>
            """
        )
    return "\n".join(cards)


def generate_dashboard_html(
    output_txt: str = "output.txt",
    insights_dir: str = "insights",
    tlp_dir: str = "TLPs",
    output_html: str = "mrsi_dashboard.html",
) -> Path:
    """
    Build a standalone multi-page HTML dashboard from generated pipeline artifacts.
    """
    digest_sections = parse_digest_sections(output_txt)
    tlp_payloads = load_tlp_payloads()
    insight_files = sorted(Path(insights_dir).glob("*.docx"))
    tlp_files = sorted(Path(tlp_dir).glob("*.docx"))
    posting_windows = extract_posting_windows(tlp_payloads)

    total_pieces = sum(len(payload.get("pieces", [])) for payload in tlp_payloads.values())
    topics_active = len(digest_sections)
    latest_date = next(iter(tlp_payloads.values()), {}).get("_date", datetime.now().strftime("%B %d, %Y"))

    dashboard_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>MRSI Dashboard</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=Plus+Jakarta+Sans:wght@600;700;800&display=swap" rel="stylesheet">
  <style>
    * {{ box-sizing: border-box; }}
    :root {{
      --bg: #f5f0fc;
      --panel: #ffffff;
      --panel-soft: #f8f5ff;
      --text: #1a1033;
      --muted: #625b7b;
      --border: #e4dbf3;
      --deep: #2d0a4e;
      --purple: #6b3fa0;
      --purple-2: #8b5cc8;
      --lime: #c9e500;
      --lime-soft: #f2f9cb;
      --ok: #166534;
      --warn: #b45309;
      --shadow: 0 20px 60px rgba(45, 10, 78, 0.08);
      --nav-w: 250px;
    }}
    body {{
      margin: 0;
      font-family: 'Inter', sans-serif;
      background:
        radial-gradient(circle at top right, rgba(201,229,0,0.12), transparent 20%),
        linear-gradient(180deg, #fbf9ff 0%, var(--bg) 100%);
      color: var(--text);
    }}
    .shell {{ display: flex; min-height: 100vh; }}
    .sidebar {{
      width: var(--nav-w);
      background: linear-gradient(180deg, #311059 0%, #22073f 100%);
      color: #fff;
      padding: 22px 14px;
      position: fixed;
      inset: 0 auto 0 0;
      overflow-y: auto;
    }}
    .brand {{
      padding: 10px 10px 20px;
      border-bottom: 1px solid rgba(255,255,255,0.08);
      margin-bottom: 16px;
    }}
    .brand-chip {{
      display: inline-flex;
      align-items: center;
      gap: 10px;
      background: rgba(255,255,255,0.08);
      border: 1px solid rgba(255,255,255,0.12);
      border-radius: 14px;
      padding: 10px 12px;
      margin-top: 14px;
    }}
    .brand-icon {{
      width: 36px;
      height: 36px;
      border-radius: 10px;
      background: var(--lime);
      color: var(--deep);
      display: grid;
      place-items: center;
      font-family: 'Plus Jakarta Sans', sans-serif;
      font-weight: 800;
    }}
    .brand h1 {{
      font-family: 'Plus Jakarta Sans', sans-serif;
      font-size: 18px;
      margin: 0;
    }}
    .brand p {{
      margin: 4px 0 0;
      color: rgba(255,255,255,0.55);
      font-size: 12px;
    }}
    .nav-title {{
      margin: 18px 8px 8px;
      font-size: 11px;
      text-transform: uppercase;
      letter-spacing: 0.12em;
      color: rgba(255,255,255,0.4);
    }}
    .nav-item {{
      display: flex;
      align-items: center;
      gap: 10px;
      padding: 11px 12px;
      border-radius: 12px;
      color: rgba(255,255,255,0.74);
      cursor: pointer;
      transition: 0.18s ease;
      margin-bottom: 4px;
    }}
    .nav-item:hover, .nav-item.active {{
      background: rgba(201,229,0,0.12);
      color: var(--lime);
    }}
    .content {{
      margin-left: var(--nav-w);
      width: calc(100% - var(--nav-w));
      padding: 28px;
    }}
    .topbar {{
      background: rgba(255,255,255,0.8);
      backdrop-filter: blur(12px);
      border: 1px solid rgba(255,255,255,0.6);
      box-shadow: var(--shadow);
      border-radius: 22px;
      padding: 18px 22px;
      display: flex;
      justify-content: space-between;
      align-items: center;
      gap: 16px;
      margin-bottom: 24px;
    }}
    .topbar h2 {{ margin: 0; font-size: 28px; font-family: 'Plus Jakarta Sans', sans-serif; }}
    .topbar p {{ margin: 5px 0 0; color: var(--muted); }}
    .btn {{
      border: 0;
      border-radius: 12px;
      background: var(--purple);
      color: #fff;
      font-weight: 700;
      padding: 12px 16px;
      cursor: pointer;
    }}
    .page {{ display: none; }}
    .page.active {{ display: block; }}
    .hero {{
      display: grid;
      grid-template-columns: 1.4fr 1fr;
      gap: 20px;
      margin-bottom: 24px;
    }}
    .hero-card, .content-card, .panel {{
      background: var(--panel);
      border: 1px solid var(--border);
      border-radius: 24px;
      box-shadow: var(--shadow);
    }}
    .hero-card {{
      padding: 24px;
      background:
        radial-gradient(circle at top right, rgba(201,229,0,0.2), transparent 25%),
        linear-gradient(135deg, rgba(107,63,160,0.08), rgba(139,92,200,0.16));
    }}
    .hero-card h3 {{
      font-family: 'Plus Jakarta Sans', sans-serif;
      font-size: 30px;
      margin: 8px 0 10px;
      max-width: 10ch;
    }}
    .hero-card p {{ color: var(--muted); max-width: 65ch; }}
    .eyebrow {{
      color: var(--purple);
      text-transform: uppercase;
      letter-spacing: 0.12em;
      font-size: 11px;
      font-weight: 700;
    }}
    .metrics {{
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 16px;
    }}
    .metric-card {{
      background: var(--panel);
      border: 1px solid var(--border);
      border-radius: 20px;
      padding: 18px;
    }}
    .metric-card.accent {{ background: linear-gradient(180deg, #ffffff, #f8fbec); }}
    .metric-label {{ color: var(--muted); font-size: 13px; }}
    .metric-value {{ font-size: 28px; font-weight: 800; margin-top: 8px; }}
    .stack {{
      display: grid;
      gap: 20px;
    }}
    .card-head {{
      display: flex;
      justify-content: space-between;
      gap: 12px;
      align-items: start;
      margin-bottom: 14px;
    }}
    .content-card {{ padding: 22px; }}
    .content-card h3 {{ margin: 6px 0; font-size: 22px; }}
    .muted {{ color: var(--muted); }}
    .doc-pill, .api-pill {{
      display: inline-flex;
      align-items: center;
      border-radius: 999px;
      padding: 7px 10px;
      font-size: 12px;
      font-weight: 700;
      background: var(--panel-soft);
      color: var(--purple);
    }}
    .doc-pill.accent {{ background: var(--lime-soft); color: #4d5b00; }}
    .clean-list {{
      padding-left: 18px;
      margin: 14px 0 18px;
      color: var(--text);
    }}
    .clean-list li {{ margin-bottom: 8px; }}
    .meta-row {{
      display: flex;
      justify-content: space-between;
      gap: 12px;
      color: var(--muted);
      font-size: 12px;
      border-top: 1px solid var(--border);
      padding-top: 14px;
    }}
    .two-col {{
      display: grid;
      grid-template-columns: 1.1fr 0.9fr;
      gap: 20px;
    }}
    .topic-list {{
      display: grid;
      gap: 12px;
      margin-top: 12px;
    }}
    .topic-item {{
      border: 1px solid var(--border);
      border-radius: 18px;
      padding: 16px;
      background: var(--panel-soft);
    }}
    .platform-grid {{
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 12px;
      margin: 18px 0;
    }}
    .platform-card {{
      border: 1px solid var(--border);
      border-radius: 18px;
      padding: 16px;
      background: #fcfbff;
    }}
    .platform-title {{ font-weight: 800; margin-bottom: 4px; }}
    .platform-angle {{ color: var(--purple); font-size: 13px; margin-bottom: 10px; }}
    .schedule-grid {{
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 16px;
    }}
    .schedule-card, .api-status {{
      border: 1px solid var(--border);
      border-radius: 18px;
      padding: 18px;
      background: #fff;
    }}
    .schedule-platform, .api-name {{ font-weight: 800; }}
    .schedule-topic, .api-key {{ color: var(--muted); font-size: 13px; margin: 5px 0 10px; }}
    .panel {{ padding: 22px; }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin-top: 12px;
    }}
    th, td {{
      text-align: left;
      padding: 12px 10px;
      border-bottom: 1px solid var(--border);
      font-size: 14px;
    }}
    th {{ color: var(--muted); font-weight: 700; }}
    .api-grid {{
      display: grid;
      grid-template-columns: repeat(3, minmax(0, 1fr));
      gap: 16px;
    }}
    .api-status.ok {{ background: linear-gradient(180deg, #fff, #f3fff3); }}
    .api-status.off {{ background: linear-gradient(180deg, #fff, #fff7f0); }}
    .empty-state {{
      background: var(--panel);
      border: 1px dashed var(--border);
      border-radius: 20px;
      padding: 24px;
      color: var(--muted);
    }}
    .settings-grid {{
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 16px;
    }}
    .setting {{
      border: 1px solid var(--border);
      border-radius: 18px;
      padding: 18px;
      background: #fff;
    }}
    .setting-label {{ color: var(--muted); font-size: 12px; text-transform: uppercase; letter-spacing: 0.08em; }}
    .setting-value {{ margin-top: 8px; font-size: 18px; font-weight: 700; }}
    @media (max-width: 1080px) {{
      .hero, .two-col, .api-grid, .schedule-grid, .platform-grid, .settings-grid {{ grid-template-columns: 1fr; }}
    }}
    @media (max-width: 820px) {{
      .sidebar {{
        position: static;
        width: 100%;
      }}
      .shell {{ display: block; }}
      .content {{
        margin-left: 0;
        width: 100%;
        padding: 18px;
      }}
      .topbar {{ flex-direction: column; align-items: start; }}
    }}
  </style>
</head>
<body>
  <div class="shell">
    <aside class="sidebar">
      <div class="brand">
        <div class="eyebrow" style="color: rgba(255,255,255,0.55);">ExoAsia</div>
        <h1>MRSI Platform</h1>
        <p>Generated pages from the latest pipeline artifacts.</p>
        <div class="brand-chip">
          <div class="brand-icon">MR</div>
          <div>
            <div style="font-weight:700;">Research Ops</div>
            <div style="font-size:12px;color:rgba(255,255,255,0.55);">{html_text(latest_date)}</div>
          </div>
        </div>
      </div>
      <div class="nav-title">Workspace</div>
      <div class="nav-item active" onclick="nav(this, 'dashboard')">Dashboard</div>
      <div class="nav-item" onclick="nav(this, 'insights')">Daily Insights</div>
      <div class="nav-item" onclick="nav(this, 'tlp')">TLP & Newsletter</div>
      <div class="nav-item" onclick="nav(this, 'autopost')">Auto-post</div>
      <div class="nav-item" onclick="nav(this, 'schedule')">Schedule</div>
      <div class="nav-item" onclick="nav(this, 'apikeys')">API Keys</div>
      <div class="nav-item" onclick="nav(this, 'settings')">Brand Settings</div>
    </aside>

    <main class="content">
      <div class="topbar">
        <div>
          <h2 id="page-title">Dashboard</h2>
          <p>Standalone MRSI control center generated by <code>main.py</code>.</p>
        </div>
        <button class="btn" onclick="window.scrollTo({{top:0,behavior:'smooth'}})">Back to Top</button>
      </div>

      <section class="page active" id="page-dashboard">
        <div class="hero">
          <div class="hero-card">
            <div class="eyebrow">Latest Run</div>
            <h3>MRSI pages are now generated from your real pipeline outputs.</h3>
            <p>This dashboard blends the original visual direction with live briefing summaries, Thought Leadership packs, publishing notes, and artifact status from the repo.</p>
          </div>
          <div class="metrics">
            {render_metric_card("Topics active", str(topics_active))}
            {render_metric_card("Insight docs", str(len(insight_files)))}
            {render_metric_card("TLP docs", str(len(tlp_files)), "accent")}
            {render_metric_card("Social pieces", str(total_pieces))}
          </div>
        </div>

        <div class="two-col">
          <div class="panel">
            <div class="card-head">
              <div>
                <div class="eyebrow">Topic Coverage</div>
                <h3>Executive summary snapshots</h3>
              </div>
            </div>
            <div class="topic-list">
              {''.join(
                  f"<div class='topic-item'><div class='platform-title'>{html_text(TOPIC_LABELS.get(key, key.title()))}</div><p class='muted'>{html_text(section.get('title', 'No title available'))}</p><p>{html_text((section.get('summary') or ['No summary available'])[0])}</p></div>"
                  for key, section in digest_sections.items()
              ) or "<div class='empty-state'>Run the pipeline once to populate topic snapshots.</div>"}
            </div>
          </div>
          <div class="panel">
            <div class="card-head">
              <div>
                <div class="eyebrow">Pipeline Artifacts</div>
                <h3>Output health</h3>
              </div>
            </div>
            <div class="topic-list">
              <div class="topic-item"><div class="platform-title">output.txt</div><p class="muted">{html_text('Present' if os.path.exists(output_txt) else 'Missing')}</p></div>
              <div class="topic-item"><div class="platform-title">Daily Insights</div><p class="muted">{html_text(f'{len(insight_files)} DOCX files ready')}</p></div>
              <div class="topic-item"><div class="platform-title">Thought Leadership</div><p class="muted">{html_text(f'{len(tlp_payloads)} JSON payloads and {len(tlp_files)} DOCX files ready')}</p></div>
              <div class="topic-item"><div class="platform-title">Dashboard HTML</div><p class="muted">{html_text(output_html)}</p></div>
            </div>
          </div>
        </div>
      </section>

      <section class="page" id="page-insights">
        <div class="stack">
          {render_insight_cards(digest_sections, insight_files)}
        </div>
      </section>

      <section class="page" id="page-tlp">
        <div class="stack">
          {render_tlp_cards(tlp_payloads, tlp_files)}
        </div>
      </section>

      <section class="page" id="page-autopost">
        <div class="panel">
          <div class="card-head">
            <div>
              <div class="eyebrow">Publishing Guidance</div>
              <h3>Platform-specific posting windows</h3>
            </div>
          </div>
          <div class="schedule-grid">
            {render_autopost_page(posting_windows)}
          </div>
        </div>
      </section>

      <section class="page" id="page-schedule">
        <div class="panel">
          <div class="card-head">
            <div>
              <div class="eyebrow">Run Schedule</div>
              <h3>Topic and asset generation table</h3>
            </div>
          </div>
          <table>
            <thead>
              <tr>
                <th>Topic</th>
                <th>Briefing</th>
                <th>Sources</th>
                <th>Expected social pieces</th>
              </tr>
            </thead>
            <tbody>
              {render_schedule_page(digest_sections, insight_files, tlp_files)}
            </tbody>
          </table>
        </div>
      </section>

      <section class="page" id="page-apikeys">
        <div class="panel">
          <div class="card-head">
            <div>
              <div class="eyebrow">Environment Check</div>
              <h3>API key availability</h3>
            </div>
          </div>
          <div class="api-grid">
            {render_api_status()}
          </div>
        </div>
      </section>

      <section class="page" id="page-settings">
        <div class="panel">
          <div class="card-head">
            <div>
              <div class="eyebrow">Brand Settings</div>
              <h3>Repo-driven configuration snapshot</h3>
            </div>
          </div>
          <div class="settings-grid">
            <div class="setting">
              <div class="setting-label">Brand</div>
              <div class="setting-value">ExoAsia MRSI</div>
            </div>
            <div class="setting">
              <div class="setting-label">Feeds configured</div>
              <div class="setting-value">{len(json.load(open('feeds.json', 'r', encoding='utf-8')))}</div>
            </div>
            <div class="setting">
              <div class="setting-label">Insight output folder</div>
              <div class="setting-value">{html_text(str(Path(insights_dir).resolve()))}</div>
            </div>
            <div class="setting">
              <div class="setting-label">TLP output folder</div>
              <div class="setting-value">{html_text(str(Path(tlp_dir).resolve()))}</div>
            </div>
          </div>
        </div>
      </section>
    </main>
  </div>

  <script>
    const pageTitles = {{
      dashboard: 'Dashboard',
      insights: 'Daily Insights',
      tlp: 'TLP & Newsletter',
      autopost: 'Auto-post',
      schedule: 'Schedule',
      apikeys: 'API Keys',
      settings: 'Brand Settings'
    }};

    function nav(el, id) {{
      document.querySelectorAll('.nav-item').forEach(item => item.classList.remove('active'));
      document.querySelectorAll('.page').forEach(page => page.classList.remove('active'));
      el.classList.add('active');
      document.getElementById('page-' + id).classList.add('active');
      document.getElementById('page-title').textContent = pageTitles[id] || id;
      window.scrollTo({{ top: 0, behavior: 'smooth' }});
    }}
  </script>
</body>
</html>
"""

    output_path = Path(output_html)
    output_path.write_text(dashboard_html, encoding="utf-8")
    return output_path


def main():
    """
    Main orchestration function for the complete pipeline.

    Order of execution:
      1. Scrape RSS feeds and enrich articles
      2. Generate daily digest via Gemini
      3. Save output.txt and Daily_Tech_Briefing.docx
      4. Run formatter.py   (reads output.txt, produces Daily Insight .docx in insights/)
      5. Run tl_summarizer  (reads output.txt, produces JSON per topic)
      6. Run tl_formatter   (reads JSON files, produces TL content packs in TLPs/)
    """

    if "--dashboard-only" in sys.argv:
        dashboard_path = generate_dashboard_html()
        print(f"Dashboard generated successfully -> {dashboard_path}")
        return

    print("Loading feeds from feeds.json...")
    try:
        with open("feeds.json", "r", encoding="utf-8") as f:
            feeds = json.load(f)
    except FileNotFoundError:
        print("Error: feeds.json not found. Please create it from the feeds.json template.")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON in feeds.json: {e}")
        sys.exit(1)

    print("Loading URL cache...")
    url_cache = load_url_cache(URL_CACHE_FILE)
    print(f"Loaded {len(url_cache)} cached URLs")

    print("Fetching RSS feeds...")
    raw_items = fetch_rss_items(feeds)
    print(f"Found {len(raw_items)} items from RSS")

    print("Enriching articles with full text...")
    enriched = enrich_articles(raw_items, url_cache, cache_file=URL_CACHE_FILE)
    print(f"Cached {len(url_cache)} URLs for future runs")

    print("Grouping articles by topic...")
    articles_by_topic: Dict[str, List[dict]] = {}
    for article in enriched:
        topic = article.get("topic", "unknown")
        if topic not in articles_by_topic:
            articles_by_topic[topic] = []
        articles_by_topic[topic].append(article)

    print("Generating digest with Gemini (Google Gemini) / AI model...")
    digest = generate_daily_digest(articles_by_topic)

    # ── Step 1: Scraper + Digest + output.txt ─────────────────────────────────
    print("Saving plain text backup to output.txt...")
    output_txt = "output.txt"
    with open(output_txt, "w", encoding="utf-8") as f:
        f.write(digest)

    print("Saving formatted document to Daily_Tech_Briefing.docx...")
    save_digest_as_docx(digest, "Daily_Tech_Briefing.docx")

    print("Digest saved successfully.")
    print("Created files:")
    print("  - output.txt")
    print("  - Daily_Tech_Briefing.docx")

    # ── Daily Insight Formatter ──────────────────────────────────────────────
    run_formatter(output_txt=output_txt)

    # ── Thought Leadership Pipeline ───────────────────────────────────────────
    run_tl_pipeline(output_txt=output_txt)

    # ── Final summary ─────────────────────────────────────────────────────────
    print("\n" + "=" * 60)
    print("PIPELINE COMPLETE")
    print("=" * 60)
    print(f"  Daily briefing   -> output.txt")
    print(f"  Briefing doc     -> Daily_Tech_Briefing.docx")
    print(f"  Daily Insights   -> insights/")
    print(f"  TL content packs -> TLPs/")
    dashboard_path = generate_dashboard_html(output_txt=output_txt)
    print(f"  Dashboard HTML   -> {dashboard_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
